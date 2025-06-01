from datetime import datetime, timedelta
import os
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy
import logging
import sys
import sqlite3
from pathlib import Path

# Настройка логирования
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('generator.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


class CertificateGenerator:
    def _init_(self):
        try:
            self.db = sqlite3.connect('certificates.db', timeout=10)
            self.create_tables()
            self.template_path = Path(_file_).parent / 'templates.docx'
            self.analyze_template()
        except Exception as e:
            logger.error(f"Ошибка инициализации: {e}")
            raise

    def create_tables(self):
        try:
            self.db.execute('''
                CREATE TABLE IF NOT EXISTS certificates (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    protocol_number TEXT UNIQUE,
                    fullname TEXT,
                    workplace TEXT,
                    position TEXT,
                    qualification_group INTEGER,
                    cert_date DATE,
                    next_cert_date DATE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            self.db.commit()
            logger.info("Проверка структуры БД выполнена")
        except sqlite3.Error as e:
            logger.error(f"Ошибка работы с БД: {e}")
            raise

    def analyze_template(self):
        try:
            doc = Document(self.template_path)
            logger.info("\nАнализ шаблона:")

            markers = [
                '{{ input_proto }}', '{{  input_proto }}',
                '{{ input_place }}', '{{ input_job }}',
                '{{ input_fullname }}', '{{ input_group }}',
                '{{ input_date }}', '{{ input_nextdate }}',
                '{{ input_what }}'
            ]

            for paragraph in doc.paragraphs:
                for marker in markers:
                    if marker in paragraph.text:
                        logger.info(f"Найден маркер {marker} в параграфе: {paragraph.text[:50]}...")
                        logger.info(f"Форматирование параграфа: {[run.text for run in paragraph.runs]}")

            group_patterns = [
                '4 гр. 1000 В и выше',
                'гр. 1000 В и выше',
                '4гр. 1000 В и выше'
            ]
            for paragraph in doc.paragraphs:
                for pattern in group_patterns:
                    if pattern in paragraph.text:
                        logger.info(f"Найден текст группы допуска: {pattern}")

            date_patterns = ['«3» декабря 2024', '«3» декабря 2025']
            for paragraph in doc.paragraphs:
                for pattern in date_patterns:
                    if pattern in paragraph.text:
                        logger.info(f"Найдена дата: {pattern}")

        except Exception as e:
            logger.error(f"Ошибка анализа шаблона: {e}")

    def get_next_protocol(self):
        year = datetime.now().year
        cursor = self.db.cursor()
        cursor.execute("SELECT COUNT(*) FROM certificates WHERE protocol_number LIKE ?",
                       (f'ПР-{year}-%',))
        count = cursor.fetchone()[0] + 1
        return f"ПР-{year}-{count:04d}"

    def format_date(self, date):
        months = {
            1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
            5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
            9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
        }
        return f"«{date.day}» {months[date.month]} {date.year}"

    def get_user_input(self):
        """Получение данных от пользователя"""
        data = {}

        print("\n=== ВВОД ДАННЫХ ДЛЯ СЕРТИФИКАТА ===\n")

        try:
            while True:
                fullname = input("Введите ФИО: ").strip()
                if len(fullname.split()) >= 2:
                    data['fullname'] = fullname
                    break
                print("Ошибка! Введите полное ФИО")

            while True:
                workplace = input("Введите место работы: ").strip()
                if workplace:
                    data['workplace'] = workplace
                    break
                print("Ошибка! Введите место работы")

            while True:
                position = input("Введите должность: ").strip()
                if position:
                    data['position'] = position
                    break
                print("Ошибка! Введите должность")

            while True:
                try:
                    group = int(input("Введите группу допуска (2-5): "))
                    if 2 <= group <= 5:
                        data['qualification_group'] = group
                        break
                    print("Ошибка! Группа должна быть от 2 до 5")
                except ValueError:
                    print("Ошибка! Введите число")

            while True:
                try:
                    date_str = input("Введите дату получения (дд.мм.гггг): ")
                    cert_date = datetime.strptime(date_str, '%d.%m.%Y')
                    data['cert_date'] = cert_date
                    break
                except ValueError:
                    print("Ошибка! Используйте формат дд.мм.гггг")

            while True:
                try:
                    date_str = input("Введите дату следующей проверки (дд.мм.гггг): ")
                    next_date = datetime.strptime(date_str, '%d.%m.%Y')
                    if next_date > data['cert_date']:
                        data['next_date'] = next_date
                        break
                    print("Ошибка! Дата следующей проверки должна быть позже даты получения")
                except ValueError:
                    print("Ошибка! Используйте формат дд.мм.гггг")

            return data

        except Exception as e:
            logger.error(f"Ошибка при вводе данных: {e}")
            raise

    def replace_text_in_paragraph(self, paragraph, replacements):
        try:
            if not paragraph.runs:
                return False

            text = paragraph.text
            original_runs = list(paragraph.runs)
            need_replace = False

            for old_text in replacements.keys():
                if old_text in text:
                    need_replace = True
                    logger.debug(f"Найден маркер для замены: {old_text}")
                    break

            if not need_replace:
                return False

            # Сохраняем форматирование
            font_props = {
                'name': original_runs[0].font.name,
                'size': original_runs[0].font.size,
                'bold': original_runs[0].font.bold,
                'italic': original_runs[0].font.italic
            }

            # Выполняем замены
            for old_text, new_text in replacements.items():
                if old_text in text:
                    logger.debug(f"Замена: {old_text} -> {new_text}")
                    text = text.replace(old_text, new_text)

            # Применяем изменения
            paragraph.clear()
            run = paragraph.add_run(text)

            # Восстанавливаем форматирование
            run.font.name = font_props['name']
            if font_props['size']:
                run.font.size = font_props['size']
            run.font.bold = font_props['bold']
            run.font.italic = font_props['italic']

            return True

        except Exception as e:
            logger.error(f"Ошибка замены текста в параграфе: {e}")
            return False

    def generate_document(self, data):
        try:
            if not self.template_path.exists():
                raise FileNotFoundError(f"Шаблон не найден: {self.template_path}")

            protocol = self.get_next_protocol()
            output_path = self.template_path.parent / f"Сертификат_{data['fullname'].split()[0]}{datetime.now():%Y%m%d%H%M%S}.docx"

            template_doc = Document(str(self.template_path))
            template_doc._body._element.xml_write_encoding = 'utf-8'

            # Форматируем данные
            formatted_fullname = data['fullname'].strip().title()
            formatted_workplace = data['workplace'].strip()
            formatted_position = data['position'].strip()

            # Обновленное форматирование группы допуска
            formatted_group = f"{data['qualification_group']} гр. 1000 В "
            formatted_group += "и выше" if data['qualification_group'] >= 3 else "и ниже"

            cert_date = self.format_date(data['cert_date'])
            next_date = self.format_date(data['next_date'])

            # Обновленный набор маркеров для замены
            replacements = {
                # Основные маркеры
                '{{ input_proto }}': protocol,
                '{{  input_proto }}': protocol,
                '{{ input_fullname }}': formatted_fullname,
                '{{ input_place }}': formatted_workplace,
                '{{ input_job }}': formatted_position,
                '{{ input_group }}': formatted_group,
                '{{ input_date }}': cert_date,
                '{{ input_nextdate }}': next_date,

                # Исправленные маркеры для группы допуска (в порядке приоритета)
                '4 гр. 1000 В и вышегр. 1000 В': formatted_group,  # Полное совпадение
                '4 гр. 1000 В и выше гр. 1000 В': formatted_group,  # С пробелом
                '4 гр. 1000 В и выше': formatted_group,  # Базовый вариант
                'гр. 1000 В и выше': formatted_group,  # Без номера группы
                '4гр. 1000 В и выше': formatted_group,  # Без пробела
                '4 4 гр. 1000 В и выше': formatted_group,  # Дублирование номера
                '4 гр. 1000 В и выше 4': formatted_group,  # Номер в конце
                '4 4': str(data['qualification_group']),  # Просто номер
                '4 4 гр.': f"{data['qualification_group']} гр.",  # Номер с гр.
                '4гр.': f"{data['qualification_group']} гр.",  # Слитно
                '4 гр.': f"{data['qualification_group']} гр.",  # С пробелом

                # Маркеры для дат
                '«3» декабря 2024': cert_date,
                '«3» декабря 2025': next_date,

                # Альтернативные варианты протокола
                '{{input_proto}}': protocol,
                '{input_proto}': protocol,
                '[input_proto]': protocol,

                # Альтернативные варианты остальных полей
                '{{input_fullname}}': formatted_fullname,
                '{input_fullname}': formatted_fullname,
                '[input_fullname]': formatted_fullname,

                '{{input_place}}': formatted_workplace,
                '{input_place}': formatted_workplace,
                '[input_place]': formatted_workplace,

                '{{input_job}}': formatted_position,
                '{input_job}': formatted_position,
                '[input_job]': formatted_position,

                # Удаляем ненужные маркеры
                '{{ input_what }}': "",
                '{{input_what}}': "",
                '{input_what}': "",
            }

            logger.info("Начало замены текста в документе")
            logger.debug(f"Используемые замены: {replacements}")

            # Замена в основном тексте
            for paragraph in template_doc.paragraphs:
                if any(marker in paragraph.text for marker in replacements.keys()):
                    logger.debug(f"Обработка параграфа: {paragraph.text}")
                    if self.replace_text_in_paragraph(paragraph, replacements):
                        logger.debug(f"Параграф после замены: {paragraph.text}")

            # Замена в таблицах
            for table in template_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if any(marker in paragraph.text for marker in replacements.keys()):
                                logger.debug(f"Обработка ячейки таблицы: {paragraph.text}")
                                if self.replace_text_in_paragraph(paragraph, replacements):
                                    logger.debug(f"Ячейка после замены: {paragraph.text}")

            # Замена в колонтитулах
            for section in template_doc.sections:
                for paragraph in section.header.paragraphs:
                    if self.replace_text_in_paragraph(paragraph, replacements):
                        logger.debug("Выполнена замена в верхнем колонтитуле")
                for paragraph in section.footer.paragraphs:
                    if self.replace_text_in_paragraph(paragraph, replacements):
                        logger.debug("Выполнена замена в нижнем колонтитуле")

            logger.info("Сохранение документа")
            template_doc.save(output_path)

            # Сохраняем в БД
            self.db.execute('''
                INSERT INTO certificates 
                (protocol_number, fullname, workplace, position, qualification_group, cert_date, next_cert_date)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                protocol,
                formatted_fullname,
                formatted_workplace,
                formatted_position,
                data['qualification_group'],
                data['cert_date'].strftime('%Y-%m-%d'),
                data['next_date'].strftime('%Y-%m-%d')
            ))
            self.db.commit()
            logger.info("Данные сохранены в БД")

            return True, output_path

        except Exception as e:
            logger.error(f"Ошибка генерации документа: {e}")
            return False, None

    def _del_(self):
        if hasattr(self, 'db'):
            self.db.close()


def main():
    try:
        generator = CertificateGenerator()
        data = generator.get_user_input()
        success, file_path = generator.generate_document(data)

        if success:
            print("\n" + "=" * 50)
            print("СЕРТИФИКАТ УСПЕШНО СОЗДАН!")
            print(f"Путь к файлу: {file_path}")
            print("=" * 50)

            os.startfile(file_path)
            os.startfile(file_path.parent)
        else:
            print("\nОШИБКА: Не удалось создать сертификат")

    except Exception as e:
        print(f"\nКРИТИЧЕСКАЯ ОШИБКА: {e}")
        logger.exception("Unexpected error occurred")

    finally:
        input("\nНажмите Enter для завершения...")


if __name__ == "_main_":
    main()